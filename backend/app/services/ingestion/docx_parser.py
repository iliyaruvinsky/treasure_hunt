from docx import Document
from pathlib import Path
from typing import Dict, Any, List
from .base_parser import BaseParser


class DOCXParser(BaseParser):
    """Parser for DOCX files"""
    
    def can_parse(self, file_path: str) -> bool:
        return self.get_file_extension(file_path) == 'docx'
    
    def parse(self, file_path: str) -> Dict[str, Any]:
        """Parse DOCX file"""
        result = {
            'metadata': {},
            'data': [],
            'errors': []
        }
        
        try:
            doc = Document(file_path)
            
            # Extract text from paragraphs
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            
            # Extract tables
            tables = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                if table_data:
                    tables.append(table_data)
            
            result['metadata'] = {
                'paragraph_count': len(paragraphs),
                'table_count': len(tables)
            }
            
            result['data'] = {
                'text': '\n\n'.join(paragraphs),
                'tables': tables
            }
            
        except Exception as e:
            result['errors'].append(f"Error parsing DOCX: {str(e)}")
        
        return result

